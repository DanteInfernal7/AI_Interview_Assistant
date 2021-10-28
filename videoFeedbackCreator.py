from docx2pdf import convert
from docx import Document
from docx.shared import Inches
import pandas as pd

def interviewFeedbackGenerator(name, speechFeedback):
    emotionPercentage= []
    df = pd.read_csv(name+".csv")
    sumOfEmotions = df['Value'].sum()
    emotions = df['Value'].tolist()
    for i in emotions:
        emotionPercentage.append((i/sumOfEmotions)*100)

    document = Document()
    mainheading = document.add_heading('Interview Report', 0)
    mainheading.alignment = 1

    document.add_paragraph('Hello '+name+', thank you for using Asistentia. Below you will find a complete report of your interview.')

    document.add_heading('Emotion Analysis', level=1)
    document.add_paragraph('Below is the analysis of the emotions and expressions portrayed by you in the video:')
    emotionGraph = document.add_picture('output\\'+name+'.png', width=Inches(6.5))
    emotionGraph.alignment = 1

    graphpara = document.add_paragraph('Graph of the emotions portrayed')
    graphpara.alignment = 1

    document.add_heading('Review of your expressions', level=1)

    if (emotionPercentage[0] > 20.00):
        document.add_paragraph('You look a bit angry in your interview, please try to look happier and presentable', style='List Bullet')
    if (emotionPercentage[1] > 30.00):
        document.add_paragraph('The interviewer may not like the looks of disgust you are giving them, try to control your emotions and seem more presentable', style='List Bullet')
    if (emotionPercentage[2] > 30.00):
        document.add_paragraph('Interviews can be intimidating but fear should not bother you', style='List Bullet')
    if (emotionPercentage[4] > 20.00):
        document.add_paragraph('Don\'t be sad because sad backwards is das and das not good! Please try to look happier and presentable', style='List Bullet')
    if (emotionPercentage[6] > 70.00):
        document.add_paragraph('If you fail to express your emotions effectively, they may not fully understand you as a person, try to be more expressive', style='List Bullet')



    records = df.values.tolist()

    table = document.add_table(rows=1 ,cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Emotions'
    hdr_cells[1].text = 'Value'
    for Emotions, Value in records:
        row_cells = table.add_row().cells
        row_cells[0].text = Emotions
        row_cells[1].text = str(Value)

    if speechFeedback is not None:
        document.add_heading('Speech Analysis', level=1)
        document.add_paragraph('Following are a few words and phrases that we noticed you use in your interview, kindly try to avoid them as they have an adverse impact on your image:')
        for speech in speechFeedback:
            document.add_paragraph(speechFeedback[speech], style='List Bullet')

    document.add_paragraph('Please feel free to check out our blog page for interview tips and tricks!')
    document.save(name+'.docx')