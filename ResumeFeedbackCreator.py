from docx2pdf import convert
from docx import Document

def resumeFeedbackGenerator(name2, resumeFeedback):
    resumeFeedback = list(resumeFeedback)
    document = Document()
    mainheading = document.add_heading('Interview Report', 0)
    mainheading.alignment = 1

    document.add_paragraph('Hello '+name2+', thank you for using Asistentia. Below you will find a complete report of your resume.')

    if resumeFeedback is not None:
        document.add_heading('Text Analysis', level=1)
        document.add_paragraph('Following are a few words and phrases that we noticed you use in your interview, kindly try to avoid them as they have an adverse impact on your image:')
        for speech in range(1,len(resumeFeedback)):
            document.add_paragraph(resumeFeedback[speech], style='List Bullet')

    document.add_paragraph('Please feel free to check out our blog page for interview tips and tricks!')
    document.save(name2+'.docx')